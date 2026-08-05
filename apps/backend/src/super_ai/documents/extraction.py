"""Indexable text extraction for supported knowledge documents."""

from __future__ import annotations

from io import BytesIO
from pathlib import PurePosixPath

from pypdf import PdfReader
from pypdf.errors import PdfReadError


def extract_indexable_text(filename: str, content: bytes) -> str:
    """Extract text from the supported upload formats before indexing."""
    suffix = PurePosixPath(filename).suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf_text(content)
    if suffix == ".docx":
        return _extract_docx_text(content)
    return _extract_markdown_text(content)


def _extract_markdown_text(content: bytes) -> str:
    try:
        text = content.decode("utf-8")
    except UnicodeDecodeError as exc:
        raise ValueError("Markdown 文件必须使用 UTF-8 编码。") from exc
    normalized = text.strip()
    if not normalized:
        raise ValueError("Markdown 文件没有可索引文本，请上传包含正文的 .md 文件。")
    return normalized


def _extract_pdf_text(content: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(content))
    except (PdfReadError, ValueError, OSError) as exc:
        raise ValueError("PDF 文件无法读取，请上传有效的 .pdf 文件。") from exc
    pages: list[str] = []
    for page in reader.pages:
        try:
            page_text = page.extract_text() or ""
        except (PdfReadError, ValueError, KeyError):
            page_text = ""
        if page_text.strip():
            pages.append(page_text.strip())
    text = "\n\n".join(pages).strip()
    if not text:
        raise ValueError("该 PDF 没有可索引文本，请上传包含可选择文本的 PDF 或转换为 Markdown。")
    return text


def _extract_docx_text(content: bytes) -> str:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise RuntimeError(
            "处理 Word 文档需要安装 python-docx，请运行 uv sync。"
        ) from exc
    try:
        document = DocxDocument(BytesIO(content))
    except (ValueError, OSError, KeyError) as exc:
        raise ValueError("Word 文件无法读取，请上传有效的 .docx 文件。") from exc
    paragraphs: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            paragraphs.append(text)
    # Also extract text from tables
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                paragraphs.append(row_text)
    text = "\n\n".join(paragraphs).strip()
    if not text:
        raise ValueError("该 Word 文档没有可索引文本，请上传包含正文的 .docx 文件。")
    return text
